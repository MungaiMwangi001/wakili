"""Text Extractor – extracts plain text from PDF, DOCX, and TXT files."""
import chardet
import structlog
import PyPDF2
from pathlib import Path
from docx import Document

log = structlog.get_logger()


def extract_text(file_path: str, file_type: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if file_type == "pdf":
        return _extract_pdf(str(path))
    elif file_type == "docx":
        return _extract_docx(str(path))
    elif file_type in ("txt",):
        return _extract_txt(str(path))
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _extract_pdf(path: str) -> str:
    try:
        # Try pdfplumber first (better extraction)
        import pdfplumber
        text_parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n".join(text_parts).strip()
    except ImportError:
        # Fallback to PyPDF2 if pdfplumber not installed
        try:
            text_parts = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n".join(text_parts).strip()
        except Exception as e:
            log.error("PDF extraction failed with PyPDF2", path=path, error=str(e))
            raise
    except Exception as e:
        log.error("PDF extraction failed with pdfplumber", path=path, error=str(e))
        # Try PyPDF2 as fallback
        try:
            text_parts = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            return "\n".join(text_parts).strip()
        except Exception as e2:
            log.error("PDF extraction failed with both methods", path=path, error=str(e2))
            raise


def _extract_docx(path: str) -> str:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n\n".join(paragraphs)


def _extract_txt(path: str) -> str:
    raw = Path(path).read_bytes()
    detected = chardet.detect(raw)
    encoding = detected.get("encoding", "utf-8") or "utf-8"
    return raw.decode(encoding, errors="replace").strip()
